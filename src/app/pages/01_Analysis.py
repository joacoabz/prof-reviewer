import streamlit as st
import tempfile
from pathlib import Path
import json
import time
import plotly.graph_objects as go
from PIL import Image
from datetime import datetime
from weasyprint import HTML
from io import BytesIO
from typing import Union

# For PDF generation

# For Word document generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Import the pipeline components
from src.pipeline.main_pipe import Pipeline
from src.openai.client import OpenAIClient

# Page configuration
st.set_page_config(page_title="Prof Reviewer - Analysis", page_icon="📝", layout="wide")

st.title("Proficiency Writing Analysis")

# Initialize session state variables if they don't exist
if "analysis_complete" not in st.session_state:
    st.session_state.analysis_complete = False

if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

if "task_description" not in st.session_state:
    st.session_state.task_description = ""

if "analysis_results" not in st.session_state:
    st.session_state.analysis_results = None

if "extracted_text" not in st.session_state:
    st.session_state.extracted_text = ""

if "processing_stage" not in st.session_state:
    st.session_state.processing_stage = ""


# Function to reset the analysis
def reset_analysis():
    st.session_state.analysis_complete = False
    st.session_state.uploaded_files = []
    st.session_state.task_description = ""
    st.session_state.analysis_results = None
    st.session_state.extracted_text = ""
    st.session_state.processing_stage = ""
    st.rerun()


# Helper function to create radar chart
def create_radar_chart(criterion_scores):
    categories = list(criterion_scores.keys())
    scores = [score for score, _ in criterion_scores.values()]

    fig = go.Figure()

    # Add radar chart
    fig.add_trace(
        go.Scatterpolar(
            r=scores,
            theta=categories,
            fill="toself",
            name="Student Score",
            line_color="rgba(31, 119, 180, 0.8)",
            fillcolor="rgba(31, 119, 180, 0.3)",
        )
    )

    # Add maximum possible score for reference
    fig.add_trace(
        go.Scatterpolar(
            r=[5, 5, 5, 5],
            theta=categories,
            fill="toself",
            name="Maximum Score",
            line=dict(color="rgba(204, 204, 204, 0.8)"),
            fillcolor="rgba(204, 204, 204, 0.1)",
        )
    )

    # Configure layout
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 5])),
        showlegend=True,
        height=400,
        margin=dict(l=80, r=80, t=20, b=20),
    )

    return fig


# Helper function to create bar chart
def create_bar_chart(criterion_scores):
    categories = list(criterion_scores.keys())
    scores = [score for score, _ in criterion_scores.values()]

    # Create color scale based on scores
    colors = [
        "#ff9999" if s < 3 else "#99ff99" if s >= 4 else "#ffcc99" for s in scores
    ]

    fig = go.Figure(
        go.Bar(
            x=scores,
            y=categories,
            orientation="h",
            marker_color=colors,
            text=scores,
            textposition="auto",
        )
    )

    fig.update_layout(
        xaxis=dict(title="Score", range=[0, 5], tickvals=[0, 1, 2, 3, 4, 5]),
        yaxis=dict(title="Criterion"),
        height=300,
        margin=dict(l=20, r=20, t=20, b=20),
    )

    return fig


# Function to run the analysis pipeline
def run_analysis(task, image_files):
    # Update processing stage
    st.session_state.processing_stage = "Initializing analysis..."

    # Create temporary directory for image files
    with tempfile.TemporaryDirectory() as temp_dir:
        # Save uploaded files to the temporary directory
        image_paths = []
        for idx, uploaded_file in enumerate(image_files):
            file_path = (
                Path(temp_dir) / f"image_{idx}.{uploaded_file.name.split('.')[-1]}"
            )
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            image_paths.append(file_path)

        try:
            # Initialize the pipeline
            pipeline = Pipeline(task=task)

            # Update processing stage
            st.session_state.processing_stage = "Extracting text from images..."
            time.sleep(0.5)  # Simulate processing time

            # Initialize OpenAI client and extract text
            openai_client = OpenAIClient()

            # Extract text from images
            st.session_state.extracted_text = pipeline.understand_solution(
                image_paths=image_paths, openai_client=openai_client
            )

            # Update processing stage
            st.session_state.processing_stage = "Analyzing task requirements..."
            time.sleep(0.5)  # Simulate processing time

            # Process task understanding
            task_understanding_obj = pipeline.task_understanding(
                task=task, openai_client=openai_client
            )
            task_understanding = json.dumps(task_understanding_obj)

            # Update processing stage
            st.session_state.processing_stage = "Assessing student solution..."
            time.sleep(0.5)  # Simulate processing time

            # Run assessment
            general_comment, criterion_scores, analysis = pipeline.assessment(
                task_understanding=task_understanding,
                students_solution=st.session_state.extracted_text,
                openai_client=openai_client,
            )

            # Update processing stage
            st.session_state.processing_stage = "Generating detailed feedback..."
            time.sleep(0.5)  # Simulate processing time

            # Generate detailed analysis
            detailed_analysis = pipeline.detailed_analysis(
                task_understanding=task_understanding,
                students_solution=st.session_state.extracted_text,
                analysis=json.dumps(analysis),
                openai_client=openai_client,
            )

            # Build full analysis for encouraging comment
            full_analysis = ""
            for analysis_item in detailed_analysis:
                analysis_str = json.dumps(analysis_item)
                full_analysis += analysis_str + "\n\n"

            # Update processing stage
            st.session_state.processing_stage = "Generating encouraging comment..."
            time.sleep(0.5)  # Simulate processing time

            # Generate encouraging comment
            encouraging_comment = pipeline.encouraging_comment(
                detailed_analysis=full_analysis,
                openai_client=openai_client,
            )

            # Store results in session state
            st.session_state.analysis_results = {
                "general_comment": general_comment,
                "criterion_scores": criterion_scores,
                "detailed_analysis": detailed_analysis,
                "analysis": analysis,
                "task_understanding": task_understanding_obj,
                "encouraging_comment": encouraging_comment,
            }

            # Set analysis complete
            st.session_state.analysis_complete = True

            # Update processing stage
            st.session_state.processing_stage = "Analysis complete!"

        except Exception as e:
            st.error(f"An error occurred during analysis: {str(e)}")
            st.session_state.processing_stage = f"Error: {str(e)}"


# Function to save analysis to history
def save_to_history():
    if not st.session_state.analysis_complete or not st.session_state.analysis_results:
        return

    # Create history directory if it doesn't exist
    history_dir = Path("logs/history")
    history_dir.mkdir(parents=True, exist_ok=True)

    # Create a history entry
    history_entry = {
        "task_description": st.session_state.task_description,
        "general_comment": st.session_state.analysis_results["general_comment"],
        "criterion_scores": {
            criterion: [score, justification]
            for criterion, (score, justification) in st.session_state.analysis_results[
                "criterion_scores"
            ].items()
        },
        "detailed_analysis": st.session_state.analysis_results["detailed_analysis"],
        "extracted_text": st.session_state.extracted_text,
        "encouraging_comment": st.session_state.analysis_results.get(
            "encouraging_comment", ""
        ),
        "total_score": sum(
            score
            for score, _ in st.session_state.analysis_results[
                "criterion_scores"
            ].values()
        ),
        "max_score": 5 * len(st.session_state.analysis_results["criterion_scores"]),
    }

    # Save to file with timestamp
    import time

    timestamp = int(time.time())
    file_path = history_dir / f"{timestamp}_analysis.json"

    with open(file_path, "w") as f:
        json.dump(history_entry, f)

    return file_path


# Function to generate PDF report
def generate_pdf_report(analysis_results, output_path=None):
    """
    Generate a PDF report from analysis results.

    Args:
        analysis_results: Dictionary containing analysis data
        output_path: Optional path to save the PDF (if None, returns bytes)

    Returns:
        Path to the saved PDF or PDF bytes
    """
    # Create a temporary directory for assets
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)

        # 1. Extract key data
        general_comment = analysis_results["general_comment"]
        criterion_scores = analysis_results["criterion_scores"]
        detailed_analysis = analysis_results["detailed_analysis"]
        encouraging_comment = analysis_results.get("encouraging_comment", "")
        extracted_text = st.session_state.extracted_text

        # Calculate total score
        total_score = sum(score for score, _ in criterion_scores.values())
        max_possible = 5 * len(criterion_scores)
        score_percentage = (total_score / max_possible) * 100

        # 2. Generate charts as images
        # Radar chart
        radar_chart = create_radar_chart(criterion_scores)
        radar_path = temp_dir / "radar_chart.png"
        radar_chart.write_image(str(radar_path), scale=3, width=800, height=600)

        # Bar chart
        bar_chart = create_bar_chart(criterion_scores)
        bar_path = temp_dir / "bar_chart.png"
        bar_chart.write_image(str(bar_path), scale=3, width=800, height=400)

        # 3. Create HTML content
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Prof Reviewer Analysis</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .section {{ margin-bottom: 30px; }}
                .chart-container {{ text-align: center; margin: 20px 0; }}
                .info-box {{ background-color: #e6f7ff; padding: 15px; border-radius: 5px; }}
                .criterion {{ margin-bottom: 15px; }}
                .criterion-header {{ font-weight: bold; background-color: #f0f0f0; padding: 8px; }}
                .criterion-content {{ padding: 8px; }}
                .issue {{ margin-bottom: 20px; border-left: 3px solid #ff6666; padding-left: 15px; }}
                .text-reference {{ background-color: #ffeeee; padding: 10px; border-radius: 5px; margin-bottom: 10px; }}
                .suggestion {{ margin-left: 20px; color: #0066cc; }}
                table {{ width: 100%; border-collapse: collapse; }}
                th, td {{ padding: 8px; text-align: left; border-bottom: 1px solid #ddd; }}
                th {{ background-color: #f2f2f2; }}
                .score {{ font-weight: bold; }}
                .pass {{ color: green; }}
                .fail {{ color: red; }}
                @page {{ size: A4; margin: 2cm; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Proficiency Writing Analysis</h1>
                <p>Generated on {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
            </div>
            
            <div class="section">
                <h2>Overall Assessment</h2>
                <div class="info-box">
                    {general_comment}
                </div>
                
                <div style="display: flex; justify-content: space-between; margin-top: 20px;">
                    <div style="flex: 1;">
                        <h3>Total Score</h3>
                        <p class="score">{total_score}/{max_possible} ({score_percentage:.1f}%)</p>
                        <p class="{"pass" if score_percentage >= 60 else "fail"}">
                            {"PASS" if score_percentage >= 60 else "NEEDS IMPROVEMENT"}
                        </p>
                    </div>
                    <div style="flex: 2;">
                        <div class="chart-container">
                            <img src="file://{radar_path}" alt="Radar Chart" style="max-width: 100%;">
                        </div>
                    </div>
                </div>
            </div>
            
            <div class="section">
                <h2>Encouraging Feedback</h2>
                <div class="info-box" style="background-color: #e6ffe6; border-left: 4px solid #66cc66;">
                    {encouraging_comment}
                </div>
            </div>
            
            <div class="section">
                <h2>Criteria Breakdown</h2>
                <div class="chart-container">
                    <img src="file://{bar_path}" alt="Criteria Scores" style="max-width: 100%;">
                </div>
                
                <table>
                    <tr>
                        <th>Criterion</th>
                        <th>Score</th>
                        <th>Justification</th>
                    </tr>
                    {"".join(f"<tr><td>{criterion.title()}</td><td>{score}/5</td><td>{justification}</td></tr>" for criterion, (score, justification) in criterion_scores.items())}
                </table>
            </div>
            
            <div class="section">
                <h2>Detailed Analysis</h2>
        """

        # Add detailed analysis grouped by category
        analysis_by_category = {}
        for item in detailed_analysis:
            category = item.get("category", "Other")
            if category not in analysis_by_category:
                analysis_by_category[category] = []
            analysis_by_category[category].append(item)

        for category, items in analysis_by_category.items():
            html_content += f"""
                <div class="criterion">
                    <div class="criterion-header">{category} Issues ({len(items)})</div>
                    <div class="criterion-content">
            """

            for i, item in enumerate(items):
                html_content += f"""
                        <div class="issue">
                            <h4>Issue {i + 1}</h4>
                            <p><strong>Text Reference:</strong></p>
                            <div class="text-reference">"{item.get("text_reference", "No text reference")}"</div>
                            <p><strong>Issue:</strong> {item.get("issue", "No issue description")}</p>
                            <p><strong>Suggestions:</strong></p>
                            <ul>
                                {"".join(f'<li class="suggestion">{suggestion}</li>' for suggestion in item.get("suggestions", []))}
                            </ul>
                        </div>
                """

            html_content += """
                    </div>
                </div>
            """

        html_content += f"""
            </div>
            
            <div class="section">
                <h2>Student Solution</h2>
                <pre style="white-space: pre-wrap; font-family: monospace; background-color: #f8f8f8; padding: 15px; border-radius: 5px;">{extracted_text}</pre>
            </div>
        </body>
        </html>
        """

        # Save HTML to temp file
        html_path = temp_dir / "report.html"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        # 4. Convert HTML to PDF
        pdf_path = output_path or (temp_dir / "analysis_report.pdf")
        HTML(string=html_content, base_url=str(temp_dir)).write_pdf(str(pdf_path))

        if output_path:
            return pdf_path
        else:
            # Return PDF as bytes if no output path provided
            with open(pdf_path, "rb") as f:
                return f.read()


# Function to generate Word document report
def generate_word_report(analysis_results, output_path=None) -> Union[Path, bytes]:
    """
    Generate a Word document report from analysis results.

    Args:
        analysis_results: Dictionary containing analysis data
        output_path: Optional path to save the Word doc (if None, returns bytes)

    Returns:
        Path to the saved document or document as bytes
    """
    # Create a temporary directory for assets if needed
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir = Path(temp_dir)

        # 1. Extract key data
        general_comment = analysis_results["general_comment"]
        criterion_scores = analysis_results["criterion_scores"]
        detailed_analysis = analysis_results["detailed_analysis"]
        encouraging_comment = analysis_results.get("encouraging_comment", "")
        extracted_text = st.session_state.extracted_text

        # Calculate total score
        total_score = sum(score for score, _ in criterion_scores.values())
        max_possible = 5 * len(criterion_scores)
        score_percentage = (total_score / max_possible) * 100

        # 2. Initialize document
        doc = Document()

        # Set document properties
        doc.core_properties.title = "Proficiency Writing Analysis"
        doc.core_properties.author = "Prof Reviewer"

        # Configure document styles - using try/except to handle potential attribute errors
        styles = doc.styles

        # Apply heading styles with error handling
        try:
            # Modify heading styles for more compact appearance
            heading1_style = styles["Heading 1"]
            heading1_style.font.size = Pt(14)
            heading1_style.font.bold = True
            heading1_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

            heading2_style = styles["Heading 2"]
            heading2_style.font.size = Pt(12)
            heading2_style.font.bold = True
            heading2_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

            # Create a style for score tables
            if "Table Content" not in styles:
                table_style = styles.add_style("Table Content", WD_STYLE_TYPE.PARAGRAPH)
                table_style.font.size = Pt(9)

            # Create a style for issue text references
            if "Text Reference" not in styles:
                ref_style = styles.add_style("Text Reference", WD_STYLE_TYPE.PARAGRAPH)
                ref_style.font.italic = True
                ref_style.font.size = Pt(9)
        except AttributeError:
            # If style attributes aren't accessible, continue without styling
            pass

        # Set narrower margins for more content
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # 4. Add General Assessment section
        doc.add_heading("Overall Assessment", level=1)
        doc.add_paragraph(general_comment)

        # Add score information
        score_para = doc.add_paragraph()
        score_para.add_run("Total Score: ").bold = True
        score_para.add_run(f"{total_score}/{max_possible} ({score_percentage:.1f}%)")

        # Add pass/fail indicator
        result_para = doc.add_paragraph()
        if score_percentage >= 60:
            result_run = result_para.add_run("PASS")
            result_run.bold = True
            result_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        else:
            result_run = result_para.add_run("NEEDS IMPROVEMENT")
            result_run.bold = True
            result_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

        # 5. Add Encouraging Feedback section
        doc.add_heading("Encouraging Feedback", level=1)
        doc.add_paragraph(encouraging_comment)

        # 6. Add Criteria Breakdown section
        doc.add_heading("Criteria Breakdown", level=1)

        # Create criteria table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Criterion"
        header_cells[1].text = "Score"
        header_cells[2].text = "Justification"

        # Style header row
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths - make criterion and score columns narrower, justification wider
        table.columns[0].width = Inches(1.2)  # Narrower column for Criterion
        table.columns[1].width = Inches(0.8)  # Narrowest column for Score
        table.columns[2].width = Inches(4.0)  # Wider column for Justification

        # Add criteria rows
        for criterion, (score, justification) in criterion_scores.items():
            row_cells = table.add_row().cells
            row_cells[0].text = criterion.title()
            row_cells[1].text = f"{score}/5"
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].text = justification

            # Apply compact styling to table content
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = "Table Content"

        # 7. Add Detailed Analysis section
        doc.add_heading("Detailed Analysis", level=1)

        # Group analysis by category
        analysis_by_category = {}
        for item in detailed_analysis:
            category = item.get("category", "Other")
            if category not in analysis_by_category:
                analysis_by_category[category] = []
            analysis_by_category[category].append(item)

        # Add each category and its issues
        for category, items in analysis_by_category.items():
            doc.add_heading(f"{category} Issues ({len(items)})", level=2)

            # Add each issue
            for i, item in enumerate(items):
                # Issue header
                issue_para = doc.add_paragraph()
                issue_para.add_run(f"Issue {i + 1}:").bold = True

                # Text reference
                doc.add_paragraph("Text Reference:", style="Heading 3")
                ref_para = doc.add_paragraph(
                    f'"{item.get("text_reference", "No text reference")}"',
                    style="Text Reference",
                )

                # Issue description
                doc.add_paragraph("Issue:", style="Heading 3")
                doc.add_paragraph(item.get("issue", "No issue description"))

                # Suggestions
                doc.add_paragraph("Suggestions:", style="Heading 3")
                suggestions = item.get("suggestions", [])
                for suggestion in suggestions:
                    suggestion_para = doc.add_paragraph()
                    suggestion_para.paragraph_format.left_indent = Inches(0.25)
                    suggestion_para.add_run("• ").bold = True
                    suggestion_para.add_run(suggestion)

                # Add separator between issues (if not the last one)
                if i < len(items) - 1:
                    doc.add_paragraph()

        # 8. Add Student Solution section
        doc.add_heading("Student Solution", level=1)
        doc.add_paragraph(extracted_text)

        # 9. Save document
        if output_path:
            doc.save(str(output_path))
            return output_path
        else:
            # Return as bytes for download
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            return bio.getvalue()


# Create a two-column layout for the main interface
col1, col2 = st.columns([1, 2])

# LEFT COLUMN: Input Section
with col1:
    st.header("Input")

    # Task Description Input
    st.subheader("Task Description")
    task_description = st.text_area(
        "Enter the task description/prompt for the writing exercise:",
        value=st.session_state.task_description,
        height=150,
        help="Paste the exact task description from the CPE exam.",
    )

    # Example task selector (optional enhancement)
    with st.expander("Need an example task?"):
        example_tasks = {
            "Essay": "Write an essay discussing the impact of technology on education, reflecting on both advantages and challenges in modern learning environments.",
            "Letter": "Write a letter to a local newspaper expressing your views on a proposal to redevelop a historic building in your town into a shopping center.",
            "Report": "Write a report for a university committee evaluating the effectiveness of online learning platforms used during the previous semester.",
        }

        task_type = st.selectbox("Select task type:", list(example_tasks.keys()))
        if st.button("Use Example"):
            st.session_state.task_description = example_tasks[task_type]
            st.rerun()

    # Image Upload Section
    st.subheader("Student Solution")
    uploaded_files = st.file_uploader(
        "Upload images of student's handwritten solution:",
        type=["jpg", "jpeg", "png", "pdf"],
        accept_multiple_files=True,
        help="You can upload multiple pages of a handwritten solution.",
    )

    # Display image preview if files are uploaded
    if uploaded_files:
        st.session_state.uploaded_files = uploaded_files
        st.write(f"Uploaded {len(uploaded_files)} image(s):")

        # Create a grid of thumbnails
        thumbnail_cols = st.columns(min(3, len(uploaded_files)))
        for idx, uploaded_file in enumerate(uploaded_files):
            col_idx = idx % 3
            with thumbnail_cols[col_idx]:
                # Display image thumbnail
                image = Image.open(uploaded_file)
                st.image(image, width=150, caption=f"Page {idx + 1}")

    # Manual text input toggle (as an alternative to OCR)
    manual_input = st.checkbox("Enter text manually instead", value=False)
    if manual_input:
        manual_text = st.text_area(
            "Enter the student's solution text:",
            height=200,
            help="Use this if you already have the text version of the student's answer.",
        )
        if manual_text:
            st.session_state.extracted_text = manual_text

    # Analysis Button
    st.subheader("Analysis Controls")

    col_analyze, col_reset = st.columns(2)

    with col_analyze:
        analyze_button = st.button(
            "Start Analysis",
            disabled=(
                not task_description
                or (
                    not uploaded_files
                    and not (manual_input and st.session_state.extracted_text)
                )
            ),
        )

    with col_reset:
        reset_button = st.button("Reset")
        if reset_button:
            reset_analysis()

    # Show processing status if analyzing
    if st.session_state.processing_stage:
        st.write("**Status:** ", st.session_state.processing_stage)

        # Display a progress bar to indicate processing
        if (
            not st.session_state.analysis_complete
            and "Error" not in st.session_state.processing_stage
        ):
            progress_bar = st.progress(0)

            # Simple simulation of progress - in a real app this would be tied to actual progress
            progress_stages = [
                "Initializing analysis...",
                "Extracting text from images...",
                "Analyzing task requirements...",
                "Assessing student solution...",
                "Generating detailed feedback...",
                "Generating encouraging comment...",
                "Analysis complete!",
            ]

            current_stage_idx = (
                progress_stages.index(st.session_state.processing_stage)
                if st.session_state.processing_stage in progress_stages
                else 0
            )
            progress_value = current_stage_idx / (len(progress_stages) - 1)
            progress_bar.progress(progress_value)

# RIGHT COLUMN: Results Display
with col2:
    # Run analysis if button is clicked
    if analyze_button:
        st.session_state.task_description = task_description

        with st.spinner("Processing..."):
            run_analysis(task_description, uploaded_files)

    # Display results if analysis is complete
    if st.session_state.analysis_complete and st.session_state.analysis_results:
        st.header("Analysis Results")

        # Create tabs for different views of the results
        tab_summary, tab_detailed, tab_text, tab_debug = st.tabs(
            ["Summary", "Detailed Analysis", "Extracted Text", "Debug Data"]
        )

        # TAB 1: SUMMARY
        with tab_summary:
            # General comment section
            st.subheader("Overall Assessment")
            st.info(st.session_state.analysis_results["general_comment"])

            # Encouraging comment section (if available)
            if "encouraging_comment" in st.session_state.analysis_results:
                st.subheader("Encouraging Feedback")
                st.success(st.session_state.analysis_results["encouraging_comment"])

            # Calculate total score
            criterion_scores = st.session_state.analysis_results["criterion_scores"]
            total_score = sum(score for score, _ in criterion_scores.values())
            max_possible = 5 * len(criterion_scores)
            score_percentage = (total_score / max_possible) * 100

            # Display total score
            col_score, col_chart = st.columns([1, 2])

            with col_score:
                st.metric(
                    "Total Score",
                    f"{total_score}/{max_possible}",
                    f"{score_percentage:.1f}%",
                )

                # Simple pass/fail indicator
                if score_percentage >= 60:
                    st.success("PASS")
                else:
                    st.error("NEEDS IMPROVEMENT")

            with col_chart:
                # Radar chart of scores
                radar_chart = create_radar_chart(criterion_scores)
                st.plotly_chart(radar_chart, use_container_width=True)

            # Criteria scorecard
            st.subheader("Criteria Breakdown")

            # Display bar chart of scores
            bar_chart = create_bar_chart(criterion_scores)
            st.plotly_chart(bar_chart, use_container_width=True)

            # Display each criterion score with justification
            for criterion, (score, justification) in criterion_scores.items():
                with st.expander(f"{criterion.title()} - Score: {score}/5"):
                    st.write(justification)

            # Action buttons
            st.subheader("Actions")
            col_export, col_save = st.columns(2)

            with col_export:
                st.write("Export Options:")

                # Replace nested columns with a horizontal layout using buttons side by side
                if st.button("Export to PDF", key="pdf_export"):
                    with st.spinner("Generating PDF report..."):
                        try:
                            # Generate a unique filename
                            timestamp = int(time.time())
                            filename = f"prof_reviewer_analysis_{timestamp}.pdf"

                            # Get PDF bytes
                            pdf_bytes = generate_pdf_report(
                                st.session_state.analysis_results
                            )

                            # Offer download via Streamlit
                            st.download_button(
                                label="Download PDF Report",
                                data=pdf_bytes,
                                file_name=filename,
                                mime="application/pdf",
                            )
                        except Exception as e:
                            st.error(f"Error generating PDF: {str(e)}")

                if st.button("Export to Word", key="word_export"):
                    with st.spinner("Generating Word report..."):
                        try:
                            # Generate a unique filename
                            timestamp = int(time.time())
                            filename = f"prof_reviewer_analysis_{timestamp}.docx"

                            # Get Word document bytes
                            docx_bytes = generate_word_report(
                                st.session_state.analysis_results
                            )

                            # Offer download via Streamlit - cast bytes explicitly
                            st.download_button(
                                label="Download Word Report",
                                data=bytes(docx_bytes)
                                if isinstance(docx_bytes, bytes)
                                else docx_bytes,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        except Exception as e:
                            st.error(f"Error generating Word document: {str(e)}")

            with col_save:
                if st.button("Save to History"):
                    saved_path = save_to_history()
                    if saved_path:
                        st.success("Analysis saved to history!")
                        # Option to navigate to history page
                        if st.button("View in History"):
                            st.switch_page("pages/02_History.py")

        # TAB 2: DETAILED ANALYSIS
        with tab_detailed:
            st.subheader("Detailed Analysis")

            # Get detailed analysis data
            detailed_analysis = st.session_state.analysis_results["detailed_analysis"]

            # Group analysis by category
            analysis_by_category = {}
            for item in detailed_analysis:
                category = item.get("category", "Other")
                if category not in analysis_by_category:
                    analysis_by_category[category] = []
                analysis_by_category[category].append(item)

            # Display analysis grouped by category
            for category, items in analysis_by_category.items():
                with st.expander(f"{category} Issues ({len(items)})"):
                    for i, item in enumerate(items):
                        st.markdown(f"### Issue {i + 1}")

                        # Text reference with highlighting
                        st.markdown("**Text Reference:**")
                        st.markdown(
                            f"<div style='background-color: #ffeeee; padding: 10px; border-radius: 5px; margin-bottom: 10px;'>\"{item.get('text_reference', 'No text reference')}\"</div>",
                            unsafe_allow_html=True,
                        )

                        # Issue description
                        st.markdown("**Issue:**")
                        st.markdown(
                            f"<div style='padding: 10px; border-left: 3px solid #ff6666; margin-bottom: 10px;'>{item.get('issue', 'No issue description')}</div>",
                            unsafe_allow_html=True,
                        )

                        # Suggestions
                        st.markdown("**Suggestions:**")
                        suggestions = item.get("suggestions", [])
                        for j, suggestion in enumerate(suggestions):
                            st.markdown(
                                f"<div style='padding: 5px 10px; margin-bottom: 5px; background-color: #e6f7ff; border-radius: 3px;'>{j + 1}. {suggestion}</div>",
                                unsafe_allow_html=True,
                            )

                        # Add separator between issues
                        if i < len(items) - 1:
                            st.markdown("---")

            # Add a section for improvement metrics
            st.subheader("Overall Improvement Areas")

            # Count issues by category
            category_counts = {
                category: len(items) for category, items in analysis_by_category.items()
            }

            # Create a bar chart of issues by category
            if category_counts:
                categories = list(category_counts.keys())
                counts = list(category_counts.values())

                fig = go.Figure(
                    go.Bar(
                        x=counts,
                        y=categories,
                        orientation="h",
                        marker_color="#ff9999",
                        text=counts,
                        textposition="auto",
                    )
                )

                fig.update_layout(
                    title="Issues by Category",
                    xaxis_title="Number of Issues",
                    yaxis_title="Category",
                    height=300,
                    margin=dict(l=20, r=20, t=40, b=20),
                )

                st.plotly_chart(fig, use_container_width=True)

                # Add a summary of total issues
                total_issues = sum(counts)
                st.markdown(f"**Total Issues Identified:** {total_issues}")

                # Find the category with the most issues
                if category_counts:
                    most_common_category = max(
                        category_counts.items(), key=lambda x: x[1]
                    )[0]
                    st.markdown(
                        f"**Most Common Issue Category:** {most_common_category}"
                    )

            # Encouraging comment section
            st.subheader("Encouraging Feedback")
            st.success(
                st.session_state.analysis_results.get(
                    "encouraging_comment", "No encouraging comment available"
                )
            )

        # TAB 3: EXTRACTED TEXT
        with tab_text:
            st.subheader("Extracted Student Solution")

            # Display the extracted text with line numbers
            st.text_area(
                "OCR Result:",
                value=st.session_state.extracted_text,
                height=400,
                disabled=True,
            )

            # Option to edit OCR errors (for future enhancement)
            if st.button("Edit Extracted Text"):
                st.session_state.show_text_editor = True

        # TAB 4: DEBUG DATA
        with tab_debug:
            st.subheader("Raw Analysis Data")

            # Task understanding
            with st.expander("Task Understanding"):
                st.json(st.session_state.analysis_results["task_understanding"])

            # Criterion scores
            with st.expander("Criterion Scores"):
                # Convert the criterion scores to a more displayable format
                score_data = {
                    criterion: {"score": score, "justification": justification}
                    for criterion, (score, justification) in criterion_scores.items()
                }
                st.json(score_data)

            # Raw analysis
            with st.expander("Raw Analysis Data"):
                st.json(st.session_state.analysis_results["analysis"])

            # Detailed analysis
            with st.expander("Detailed Analysis"):
                st.json(st.session_state.analysis_results["detailed_analysis"])

            # Encouraging comment
            with st.expander("Encouraging Comment"):
                st.write(
                    st.session_state.analysis_results.get(
                        "encouraging_comment", "No encouraging comment available"
                    )
                )
